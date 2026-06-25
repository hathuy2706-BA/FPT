#!/usr/bin/env python3
"""Generate Product Backlog Word document for Product Team."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/Users/hathuy/Documents/FPT-1/docs/product_backlog_team.docx"

# ── Color palette ────────────────────────────────────────────────────────────
ORANGE   = RGBColor(0xE8, 0x6A, 0x23)   # FPT orange
DARK     = RGBColor(0x1F, 0x27, 0x37)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_BG  = RGBColor(0xF5, 0xF5, 0xF5)
LIGHT_OR = RGBColor(0xFF, 0xF0, 0xE8)
BLUE_H   = RGBColor(0x1E, 0x4D, 0x8C)

MUST     = RGBColor(0xC0, 0x39, 0x2B)
SHOULD   = RGBColor(0xE6, 0x7E, 0x22)
COULD    = RGBColor(0x27, 0xAE, 0x60)
WONT     = RGBColor(0x95, 0xA5, 0xA6)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, edges=("top","bottom","left","right"), size=4, color="E8E8E8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBr = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBr.append(el)
    tcPr.append(tcBr)

def cell_para(cell, text, bold=False, size=11, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    return p

def add_para(doc, text, bold=False, size=13, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    return p

def moscow_color(label: str) -> RGBColor:
    m = label.strip().upper()
    if m.startswith("M"):  return MUST
    if m.startswith("S"):  return SHOULD
    if m.startswith("C"):  return COULD
    return WONT

# ── Data ─────────────────────────────────────────────────────────────────────
EPICS = [
    ("EP01", "Quản lý Task & Công việc"),
    ("EP02", "Sprint & Iteration Planning"),
    ("EP03", "Dashboard & Báo cáo tiến độ"),
    ("EP04", "Phân quyền & Quản lý thành viên"),
    ("EP05", "Tích hợp & Thông báo"),
]

# Columns: ID | Epic | User Story | AC (Gherkin) | Priority | SP | Sprint
BACKLOG = [
    # ── EP01 ──────────────────────────────────────────────────────────────
    (
        "US-001", "EP01",
        "Tạo task mới\n\nAs a BA/PO, I want to create a new task with full details (title, description, type, assignee, priority, deadline) so that the team has a clear and trackable work item.",
        "GIVEN tôi ở trang Backlog\n"
        "WHEN tôi nhấn \"Tạo Task\" và điền đầy đủ các trường bắt buộc\n"
        "THEN task được tạo và hiển thị trong backlog với status \"To Do\"\n\n"
        "GIVEN tôi để trống trường Title\n"
        "WHEN tôi nhấn \"Lưu\"\n"
        "THEN hệ thống hiển thị lỗi \"Title không được để trống\"",
        "Must Have", "3", "Sprint 1"
    ),
    (
        "US-002", "EP01",
        "Cập nhật trạng thái task (Kanban)\n\nAs a team member, I want to drag & drop tasks across columns (To Do → In Progress → Review → Done) so that the team can visualize workflow in real time.",
        "GIVEN task đang ở cột \"To Do\"\n"
        "WHEN tôi kéo task sang \"In Progress\"\n"
        "THEN task chuyển trạng thái và lưu timestamp bắt đầu\n\n"
        "GIVEN task chuyển sang \"Done\"\n"
        "THEN hệ thống ghi nhận completion date và tính cycle time",
        "Must Have", "3", "Sprint 1"
    ),
    (
        "US-003", "EP01",
        "Assign task cho thành viên\n\nAs a PO, I want to assign a task to one or more team members so that responsibilities are clear.",
        "GIVEN tôi đang xem chi tiết task\n"
        "WHEN tôi chọn assignee từ dropdown danh sách thành viên\n"
        "THEN task hiển thị avatar assignee và thành viên nhận notification\n\n"
        "GIVEN task chưa có assignee\n"
        "THEN hệ thống hiển thị cảnh báo \"Chưa assign\" trên board",
        "Must Have", "2", "Sprint 1"
    ),
    (
        "US-004", "EP01",
        "Đặt độ ưu tiên & deadline\n\nAs a PO, I want to set priority (Critical/High/Medium/Low) and deadline for each task so that the team can focus on what matters most.",
        "GIVEN tôi tạo hoặc chỉnh sửa task\n"
        "WHEN tôi chọn mức ưu tiên và nhập ngày deadline\n"
        "THEN task hiển thị badge màu tương ứng và đếm ngược đến deadline\n\n"
        "GIVEN deadline < 2 ngày\n"
        "THEN badge chuyển màu đỏ và gửi reminder notification",
        "Must Have", "2", "Sprint 1"
    ),
    (
        "US-005", "EP01",
        "Tìm kiếm & lọc task\n\nAs a team member, I want to search and filter tasks by keyword, assignee, status, sprint, or tag so that I can quickly find relevant work.",
        "GIVEN tôi nhập keyword vào ô tìm kiếm\n"
        "WHEN tôi gõ ≥ 2 ký tự\n"
        "THEN danh sách task filter realtime theo title và description\n\n"
        "GIVEN tôi kết hợp filter (status=In Progress + assignee=Tôi)\n"
        "THEN chỉ hiển thị task thỏa mãn tất cả điều kiện",
        "Must Have", "3", "Sprint 2"
    ),
    (
        "US-006", "EP01",
        "Comment & đính kèm file trên task\n\nAs a BA, I want to add comments and attach files (doc, pdf, image) to a task so that the discussion and evidence stay in context.",
        "GIVEN tôi ở trang chi tiết task\n"
        "WHEN tôi nhập comment và nhấn \"Gửi\"\n"
        "THEN comment hiển thị với timestamp và avatar người dùng\n\n"
        "GIVEN tôi đính kèm file > 25 MB\n"
        "THEN hệ thống hiển thị lỗi giới hạn kích thước file",
        "Should Have", "3", "Sprint 2"
    ),
    # ── EP02 ──────────────────────────────────────────────────────────────
    (
        "US-007", "EP02",
        "Tạo và quản lý Sprint\n\nAs a PO, I want to create sprints with name, goal, start date, and end date so that the team works in structured time-boxes.",
        "GIVEN tôi ở màn Sprint Planning\n"
        "WHEN tôi tạo Sprint với đầy đủ thông tin\n"
        "THEN Sprint xuất hiện trong danh sách với status \"Planned\"\n\n"
        "GIVEN sprint end date < start date\n"
        "THEN hệ thống báo lỗi validation",
        "Must Have", "3", "Sprint 1"
    ),
    (
        "US-008", "EP02",
        "Thêm task vào Sprint\n\nAs a PO, I want to move tasks from backlog into a sprint so that the sprint scope is clearly defined.",
        "GIVEN tôi đang xem Sprint backlog\n"
        "WHEN tôi kéo task từ Product Backlog vào Sprint\n"
        "THEN task gắn với Sprint và cập nhật tổng Story Points\n\n"
        "GIVEN sprint đã Start\n"
        "WHEN tôi thêm task mới\n"
        "THEN hiển thị cảnh báo \"Thêm task giữa Sprint — cần PO xác nhận\"",
        "Must Have", "3", "Sprint 2"
    ),
    (
        "US-009", "EP02",
        "Burn-down chart theo Sprint\n\nAs a PO, I want to see a burn-down chart for the active sprint so that I can track whether the team is on pace to complete the sprint goal.",
        "GIVEN sprint đang chạy\n"
        "WHEN tôi mở tab Burn-down\n"
        "THEN biểu đồ hiển thị đường ideal vs actual remaining SP theo ngày\n\n"
        "GIVEN tất cả task của sprint = Done\n"
        "THEN đường actual chạm 0 và hiển thị badge \"Sprint Completed\"",
        "Should Have", "5", "Sprint 3"
    ),
    (
        "US-010", "EP02",
        "Sprint Retrospective notes\n\nAs a Scrum Master/PO, I want to record retrospective notes (Went Well, Improvement, Action Items) after each sprint so that lessons learned are persisted.",
        "GIVEN sprint đã kết thúc\n"
        "WHEN tôi mở trang Retrospective\n"
        "THEN có 3 cột Went Well / Improvement / Action Items để thêm sticky notes\n\n"
        "GIVEN tôi lưu retro\n"
        "THEN dữ liệu liên kết với sprint đó và có thể xem lại trong lịch sử",
        "Should Have", "3", "Sprint 3"
    ),
    # ── EP03 ──────────────────────────────────────────────────────────────
    (
        "US-011", "EP03",
        "Dashboard cá nhân\n\nAs a team member, I want a personal dashboard showing my assigned tasks, deadlines, and today's priorities so that I can plan my workday efficiently.",
        "GIVEN tôi đăng nhập\n"
        "WHEN tôi mở My Dashboard\n"
        "THEN hiển thị: task đang làm, due today, overdue, và activity feed\n\n"
        "GIVEN không có task nào\n"
        "THEN hiển thị empty state \"Không có task nào — hãy tận hưởng ngày hôm nay!\"",
        "Must Have", "3", "Sprint 2"
    ),
    (
        "US-012", "EP03",
        "Dashboard Team — tổng quan tiến độ\n\nAs a PO, I want a team dashboard with task distribution by status, workload per member, and sprint progress so that I can manage team capacity.",
        "GIVEN tôi ở Team Dashboard\n"
        "WHEN sprint đang active\n"
        "THEN hiển thị: % Done, số task theo status, workload chart từng thành viên\n\n"
        "GIVEN thành viên có > 5 task In Progress\n"
        "THEN hiển thị cảnh báo overload với màu cam",
        "Must Have", "5", "Sprint 3"
    ),
    (
        "US-013", "EP03",
        "Báo cáo Velocity Sprint-over-Sprint\n\nAs a PO, I want to see team velocity (SP completed per sprint) across multiple sprints so that I can make reliable release forecasts.",
        "GIVEN có ≥ 2 sprint đã hoàn thành\n"
        "WHEN tôi mở Velocity Report\n"
        "THEN biểu đồ cột hiển thị SP committed vs SP completed cho từng sprint\n\n"
        "GIVEN tôi hover vào cột sprint\n"
        "THEN tooltip hiển thị danh sách task của sprint đó",
        "Could Have", "5", "Sprint 4"
    ),
    (
        "US-014", "EP03",
        "Xuất báo cáo tiến độ (PDF/Excel)\n\nAs a PO, I want to export sprint/project progress reports to PDF or Excel so that I can share with stakeholders.",
        "GIVEN tôi chọn loại báo cáo và khoảng thời gian\n"
        "WHEN tôi nhấn \"Xuất\"\n"
        "THEN file được tải về trong ≤ 5 giây, đúng template FPT\n\n"
        "GIVEN có > 500 task trong report\n"
        "THEN hệ thống xử lý background và gửi email khi xong",
        "Could Have", "5", "Sprint 4"
    ),
    # ── EP04 ──────────────────────────────────────────────────────────────
    (
        "US-015", "EP04",
        "Quản lý vai trò thành viên\n\nAs an Admin/PO, I want to assign roles (PO, BA, Dev, QA, Viewer) to team members so that each person has appropriate access and permissions.",
        "GIVEN tôi là Admin\n"
        "WHEN tôi vào Settings > Members và gán role cho user\n"
        "THEN user nhận email thông báo role mới, quyền thay đổi ngay lập tức\n\n"
        "GIVEN user bị remove khỏi project\n"
        "THEN tất cả task assign cho user đó hiển thị cảnh báo \"Unassigned\"",
        "Must Have", "3", "Sprint 1"
    ),
    (
        "US-016", "EP04",
        "Phân quyền theo Project\n\nAs an Admin, I want to control which members can view, edit, or manage each project so that sensitive projects are protected.",
        "GIVEN tôi tạo project với visibility = Private\n"
        "WHEN user không được mời\n"
        "THEN project không xuất hiện trong danh sách của user đó\n\n"
        "GIVEN PO cố xóa task của người khác\n"
        "THEN hệ thống cho phép (PO có quyền manage)\n\n"
        "GIVEN Viewer cố tạo task\n"
        "THEN hệ thống trả về lỗi 403 Forbidden",
        "Must Have", "3", "Sprint 2"
    ),
    # ── EP05 ──────────────────────────────────────────────────────────────
    (
        "US-017", "EP05",
        "Thông báo Email & In-app\n\nAs a team member, I want to receive notifications (in-app + email) when I'm assigned a task, mentioned in a comment, or when my task is due soon so that I never miss important updates.",
        "GIVEN task được assign cho tôi\n"
        "WHEN assignee thay đổi\n"
        "THEN tôi nhận in-app notification ngay và email trong vòng 1 phút\n\n"
        "GIVEN tôi tắt email notification trong Settings\n"
        "THEN chỉ nhận in-app, không gửi email",
        "Should Have", "3", "Sprint 3"
    ),
    (
        "US-018", "EP05",
        "Tích hợp Slack\n\nAs a PO, I want to receive sprint summary and task update notifications in a Slack channel so that the team stays aligned without switching tools.",
        "GIVEN Slack integration đã được cấu hình\n"
        "WHEN task chuyển sang \"Done\"\n"
        "THEN bot Slack gửi message vào channel với link task\n\n"
        "GIVEN sprint kết thúc\n"
        "THEN bot gửi Sprint Summary (SP done/total, % complete) vào channel",
        "Could Have", "5", "Sprint 4"
    ),
    (
        "US-019", "EP05",
        "Tích hợp Jira (Import/Export)\n\nAs a BA, I want to import existing tasks from Jira and export task data to Jira so that the team can migrate gradually without losing history.",
        "GIVEN tôi cung cấp Jira API token và project key\n"
        "WHEN tôi chạy Import\n"
        "THEN tất cả issues của Jira project được tạo trong hệ thống với mapping: Issue → Task, Story → User Story, Bug → Bug\n\n"
        "GIVEN import thất bại một phần\n"
        "THEN hệ thống cung cấp error log theo từng item",
        "Won't Have (v1)", "8", "Sprint 5+"
    ),
]

SPRINT_SUMMARY = [
    ("Sprint 1", "2 tuần", "US-001, 002, 003, 004, 007, 015", "19 SP", "Core task management + roles"),
    ("Sprint 2", "2 tuần", "US-005, 006, 008, 011, 016",      "14 SP", "Search, filter, sprint scope, dashboard cá nhân"),
    ("Sprint 3", "2 tuần", "US-009, 010, 012, 017",           "14 SP", "Burn-down, retro, team dashboard, notifications"),
    ("Sprint 4", "2 tuần", "US-013, 014, 018",                "15 SP", "Velocity report, export, Slack"),
    ("Sprint 5+","TBD",    "US-019",                          "8 SP",  "Jira integration (ngoài phạm vi v1)"),
]

# ── Build Document ────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    # ── Cover page ───────────────────────────────────────────────────────
    # Logo via URL
    try:
        import urllib.request, tempfile, os
        logo_url = "https://upload.wikimedia.org/wikipedia/commons/1/11/FPT_logo_2010.svg"
        # Use text fallback since SVG not directly supported
        raise Exception("use text")
    except Exception:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("FPT SOFTWARE")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = ORANGE
        run.font.name = "Times New Roman"

    doc.add_paragraph()

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PRODUCT BACKLOG")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK
    r.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Hệ thống Quản lý Task — Đội Product")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = ORANGE
    r2.font.name = "Times New Roman"

    doc.add_paragraph()

    # Meta table
    meta = doc.add_table(rows=5, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in meta.columns:
        for cell in col.cells:
            cell.width = Cm(4.5)

    def meta_row(row_idx, l1, v1, l2, v2):
        row = meta.rows[row_idx]
        set_cell_bg(row.cells[0], "E86A23"); cell_para(row.cells[0], l1, bold=True, size=11, color=WHITE)
        cell_para(row.cells[1], v1, size=11)
        set_cell_bg(row.cells[2], "E86A23"); cell_para(row.cells[2], l2, bold=True, size=11, color=WHITE)
        cell_para(row.cells[3], v2, size=11)

    meta_row(0, "Dự án:",      "Task Management System",   "Phiên bản:", "1.0")
    meta_row(1, "Team:",       "Product Team (BA/PO/Dev)",  "Ngày tạo:",  "25/06/2026")
    meta_row(2, "Author:",     "Product Owner / BA",        "Trạng thái:", "Draft")
    meta_row(3, "Tổng Epics:", "5",                          "Tổng US:",   str(len(BACKLOG)))
    meta_row(4, "Tổng SP:",    "75",                         "Sprints:",   "4 sprints (v1)")

    doc.add_paragraph()

    # ── Section 1: Epics ─────────────────────────────────────────────────
    add_para(doc, "1. EPIC OVERVIEW", bold=True, size=14, color=ORANGE, space_before=10)

    epic_tbl = doc.add_table(rows=1, cols=3)
    epic_tbl.style = "Table Grid"
    epic_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = epic_tbl.rows[0]
    for cell, txt in zip(hdr.cells, ["Epic ID", "Tên Epic", "Mô tả ngắn"]):
        set_cell_bg(cell, "1F2737")
        cell_para(cell, txt, bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    epic_descs = [
        "Tạo, cập nhật, assign, filter, comment task",
        "Tạo Sprint, scope, burn-down, retrospective",
        "Personal & team dashboard, velocity, export",
        "Vai trò, quyền truy cập theo project",
        "Email, in-app notification, Slack, Jira",
    ]
    for (eid, ename), edesc in zip(EPICS, epic_descs):
        row = epic_tbl.add_row()
        set_cell_bg(row.cells[0], "FFF0E8")
        cell_para(row.cells[0], eid,   bold=True,  size=11, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(row.cells[1], ename, bold=True,  size=11)
        cell_para(row.cells[2], edesc, bold=False, size=10)

    for col_idx, width in enumerate([Cm(2.5), Cm(6.5), Cm(9)]):
        for cell in epic_tbl.columns[col_idx].cells:
            cell.width = width

    doc.add_paragraph()

    # ── Section 2: Product Backlog ────────────────────────────────────────
    add_para(doc, "2. PRODUCT BACKLOG CHI TIẾT", bold=True, size=14, color=ORANGE, space_before=10)

    # MoSCoW legend
    legend_p = doc.add_paragraph()
    legend_p.paragraph_format.space_before = Pt(2)
    legend_p.paragraph_format.space_after  = Pt(6)
    legend_p.add_run("MoSCoW: ").bold = True
    for label, col in [("Must Have", MUST), ("Should Have", SHOULD),
                        ("Could Have", COULD), ("Won't Have (v1)", WONT)]:
        r = legend_p.add_run(f"  ● {label}  ")
        r.font.color.rgb = col
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    # Backlog table
    cols = ["ID", "Epic", "User Story", "Acceptance Criteria (Gherkin)", "Priority", "SP", "Sprint"]
    widths = [Cm(1.8), Cm(1.5), Cm(5.5), Cm(6.8), Cm(2.3), Cm(0.9), Cm(1.8)]

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = tbl.rows[0]
    for i, (cell, txt) in enumerate(zip(hdr_row.cells, cols)):
        set_cell_bg(cell, "1F2737")
        cell_para(cell, txt, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.width = widths[i]

    prev_epic = None
    for us_id, epic, story, ac, priority, sp, sprint in BACKLOG:
        row = tbl.add_row()
        row.cells[0].width = widths[0]

        # Alternating row bg
        bg = "FFFFFF" if (BACKLOG.index((us_id, epic, story, ac, priority, sp, sprint)) % 2 == 0) else "F9F9F9"
        for ci in range(len(cols)):
            row.cells[ci].width = widths[ci]
            set_cell_bg(row.cells[ci], bg)

        # ID
        cell_para(row.cells[0], us_id, bold=True, size=10, color=BLUE_H, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Epic
        epic_color = ORANGE if epic != prev_epic else RGBColor(0x80, 0x80, 0x80)
        cell_para(row.cells[1], epic, bold=(epic != prev_epic), size=10, color=epic_color, align=WD_ALIGN_PARAGRAPH.CENTER)
        prev_epic = epic

        # User Story — split title & story
        parts = story.split("\n\n", 1)
        tc = row.cells[2]
        p  = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(1)
        r_title = p.add_run(parts[0])
        r_title.bold = True
        r_title.font.size = Pt(10)
        r_title.font.name = "Times New Roman"
        r_title.font.color.rgb = DARK
        if len(parts) > 1:
            p2 = tc.add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after  = Pt(2)
            r2 = p2.add_run(parts[1])
            r2.font.size  = Pt(9)
            r2.font.name  = "Times New Roman"
            r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # AC
        cell_para(row.cells[3], ac, size=9)
        row.cells[3].paragraphs[0].paragraph_format.space_before = Pt(2)
        row.cells[3].paragraphs[0].paragraph_format.space_after  = Pt(2)

        # Priority
        mc = moscow_color(priority)
        p_cell = row.cells[4]
        p_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        p_cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
        r_prio = p_cell.paragraphs[0].add_run(priority)
        r_prio.bold = True
        r_prio.font.size = Pt(9)
        r_prio.font.name = "Times New Roman"
        r_prio.font.color.rgb = mc

        # SP
        cell_para(row.cells[5], sp, bold=True, size=11, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Sprint
        cell_para(row.cells[6], sprint, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Section 3: Sprint Planning ────────────────────────────────────────
    add_para(doc, "3. SPRINT PLANNING SƠ BỘ (Release v1)", bold=True, size=14, color=ORANGE, space_before=10)

    sp_tbl = doc.add_table(rows=1, cols=5)
    sp_tbl.style = "Table Grid"
    sp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, txt in zip(sp_tbl.rows[0].cells, ["Sprint", "Thời lượng", "User Stories", "Story Points", "Sprint Goal"]):
        set_cell_bg(cell, "1F2737")
        cell_para(cell, txt, bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    sp_widths = [Cm(2.0), Cm(2.5), Cm(5.0), Cm(2.5), Cm(6.5)]
    for i, (sprint, duration, stories, pts, goal) in enumerate(SPRINT_SUMMARY):
        row = sp_tbl.add_row()
        bg = "FFF0E8" if i % 2 == 0 else "FFFFFF"
        vals = [sprint, duration, stories, pts, goal]
        for ci, (cell, val) in enumerate(zip(row.cells, vals)):
            set_cell_bg(cell, bg)
            cell.width = sp_widths[ci]
            is_sprint = (ci == 0)
            cell_para(cell, val, bold=is_sprint, size=10,
                      color=(ORANGE if is_sprint else DARK),
                      align=(WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT))

    doc.add_paragraph()

    # ── Section 4: Definition of Done ────────────────────────────────────
    add_para(doc, "4. DEFINITION OF DONE (DoD)", bold=True, size=14, color=ORANGE, space_before=10)

    dod_items = [
        "Code đã được review bởi ít nhất 1 thành viên khác.",
        "Unit test coverage ≥ 80% cho business logic.",
        "Acceptance Criteria đã pass toàn bộ (manual hoặc automated).",
        "Không có bug P1/P2 mở tại thời điểm demo.",
        "UI đã được kiểm tra trên Chrome, Safari, Firefox (responsive).",
        "Tài liệu kỹ thuật (API doc, flow) được cập nhật nếu có thay đổi.",
        "PO/BA xác nhận nghiệm thu feature trước khi chuyển sang Done.",
    ]

    dod_tbl = doc.add_table(rows=len(dod_items), cols=2)
    dod_tbl.style = "Table Grid"
    dod_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, item in enumerate(dod_items):
        row = dod_tbl.rows[i]
        bg = "F9F9F9" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row.cells[0], "27AE60")
        cell_para(row.cells[0], "✓", bold=True, size=12, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[0].width = Cm(1.0)
        set_cell_bg(row.cells[1], bg)
        cell_para(row.cells[1], item, size=11)
        row.cells[1].width = Cm(17.5)

    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────
    note = add_para(doc, "* Tài liệu này được tạo tự động bởi FPT BA/PO AI Platform. Phiên bản 1.0 — 25/06/2026.",
                    size=9, color=RGBColor(0x88, 0x88, 0x88),
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    build()
